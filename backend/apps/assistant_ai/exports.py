from __future__ import annotations

import io
from html import escape
from pathlib import Path
from typing import Iterable

from django.utils.text import slugify
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem, PageBreak

from .models import AIDraft

BRAND = colors.HexColor("#ff641a")
NAVY = colors.HexColor("#0b1f3a")
MUTED = colors.HexColor("#64748b")


def _strings(values) -> list[str]:
    if not isinstance(values, list):
        return []
    return [str(v).strip() for v in values if str(v).strip()]


def draft_sections(draft: AIDraft) -> list[tuple[str, str | list[str]]]:
    p = draft.payload or {}
    sections: list[tuple[str, str | list[str]]] = []
    if draft.kind == AIDraft.Kind.COVER_LETTER:
        sections.append(("Lettre de motivation", str(p.get("content") or "")))
        points = _strings(p.get("key_points"))
        if points:
            sections.append(("Points clés", points))
    elif draft.kind == AIDraft.Kind.CV_IMPROVEMENT:
        try:
            profile = draft.user.candidate_profile
        except Exception:
            profile = None
        identity = [draft.user.get_full_name() or draft.user.username, draft.user.email]
        if getattr(draft.user, "country", ""):
            identity.append(str(draft.user.country))
        sections.append(("Coordonnées", " · ".join(value for value in identity if value)))
        headline = str(p.get("professional_headline") or getattr(profile, "headline", "") or "")
        summary = str(p.get("summary") or getattr(profile, "summary", "") or "")
        skills = _strings(p.get("skills")) or _strings(getattr(profile, "skills", []))
        if headline:
            sections.append(("Accroche professionnelle", headline))
        if summary:
            sections.append(("Résumé professionnel", summary))
        if skills:
            sections.append(("Compétences", skills))
        if profile:
            roles = _strings(getattr(profile, "desired_roles", []))
            if roles:
                sections.append(("Postes recherchés", roles))
            sections.append(("Expérience", f"{int(getattr(profile, 'years_experience', 0) or 0)} an(s) d'expérience déclarée"))
        for key, label in (("achievement_rewrites", "Réalisations reformulées"), ("recommendations", "Recommandations de relecture")):
            vals = _strings(p.get(key))
            if vals:
                sections.append((label, vals))
    elif draft.kind == AIDraft.Kind.LEARNING_GAP_PLAN:
        missing = _strings(p.get("missing_skills"))
        if missing:
            sections.append(("Compétences à renforcer", missing))
        actions = []
        for raw in p.get("actions") or []:
            if isinstance(raw, dict):
                skill = str(raw.get("skill") or "Compétence")
                action = str(raw.get("action") or "")
                if action:
                    actions.append(f"{skill} — {action}")
        if actions:
            sections.append(("Plan d'action", actions))
    elif draft.kind == AIDraft.Kind.INTERVIEW_PREP:
        if p.get("pitch"):
            sections.append(("Pitch", str(p.get("pitch"))))
        for key, label in (("likely_questions", "Questions probables"), ("star_examples", "Exemples STAR"), ("questions_to_ask", "Questions à poser"), ("checklist", "Checklist")):
            vals = _strings(p.get(key))
            if vals:
                sections.append((label, vals))
    elif draft.kind == AIDraft.Kind.QUIZ:
        vals = []
        for i, raw in enumerate(p.get("questions") or [], 1):
            if isinstance(raw, dict):
                vals.append(f"{i}. {raw.get('question', '')} — Réponse : {raw.get('correct_answer', '')}")
        if vals:
            sections.append(("Questions", vals))
    elif draft.kind == AIDraft.Kind.COURSE_OUTLINE:
        if p.get("description"):
            sections.append(("Description", str(p.get("description"))))
        vals = []
        for raw in p.get("sections") or []:
            if isinstance(raw, dict):
                title = str(raw.get("title") or "Section")
                lessons = ", ".join(_strings(raw.get("lessons")))
                vals.append(f"{title}: {lessons}" if lessons else title)
        if vals:
            sections.append(("Plan", vals))
    elif draft.kind == AIDraft.Kind.MENTOR_PLAN:
        for key, label in (("objectives", "Objectifs"), ("agenda", "Agenda"), ("questions", "Questions"), ("follow_up", "Suivi")):
            vals = _strings(p.get(key))
            if vals:
                sections.append((label, vals))
    elif draft.kind == AIDraft.Kind.INTERVIEW_RUBRIC:
        for key, label in (("criteria", "Critères"), ("questions", "Questions")):
            vals = _strings(p.get(key))
            if vals:
                sections.append((label, vals))
    if p.get("opportunity"):
        sections.insert(0, ("Opportunité", str(p.get("opportunity"))))
    return sections or [("Contenu", "Aucun contenu exportable.")]


def export_filename(draft: AIDraft, extension: str) -> str:
    base = slugify(draft.title)[:70] or f"brouillon-{draft.id}"
    return f"kalanpro-{base}.{extension}"


def build_docx(draft: AIDraft) -> bytes:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    title = document.add_paragraph()
    run = title.add_run("KalanPro")
    run.bold = True; run.font.size = Pt(12); run.font.color.rgb = RGBColor(255, 100, 26)
    heading = document.add_heading(draft.title, level=1)
    heading.runs[0].font.color.rgb = RGBColor(11, 31, 58)
    document.add_paragraph(f"Document généré depuis KalanPro AI · {draft.get_kind_display()}")
    for label, value in draft_sections(draft):
        document.add_heading(label, level=2)
        if isinstance(value, list):
            for item in value:
                document.add_paragraph(item, style="List Bullet")
        else:
            for paragraph in str(value).split("\n"):
                if paragraph.strip():
                    document.add_paragraph(paragraph.strip())
    document.add_paragraph("Relisez toujours un contenu généré par IA avant utilisation.")
    out = io.BytesIO()
    document.save(out)
    return out.getvalue()


def build_pdf(draft: AIDraft) -> bytes:
    out = io.BytesIO()
    doc = SimpleDocTemplate(out, pagesize=A4, rightMargin=18*mm, leftMargin=18*mm, topMargin=18*mm, bottomMargin=18*mm, title=draft.title, author="KalanPro AI")
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("KPTitle", parent=styles["Title"], textColor=NAVY, fontName="Helvetica-Bold", fontSize=20, leading=24, alignment=TA_LEFT, spaceAfter=8)
    h2 = ParagraphStyle("KPH2", parent=styles["Heading2"], textColor=NAVY, fontName="Helvetica-Bold", fontSize=12, leading=15, spaceBefore=10, spaceAfter=5)
    body = ParagraphStyle("KPBody", parent=styles["BodyText"], textColor=colors.HexColor("#334155"), fontName="Helvetica", fontSize=9.5, leading=14, spaceAfter=5)
    small = ParagraphStyle("KPSmall", parent=body, textColor=MUTED, fontSize=8, leading=11)
    story = [Paragraph("<b><font color='#ff641a'>KalanPro</font></b>", body), Paragraph(escape(draft.title), title_style), Paragraph(f"{escape(draft.get_kind_display())} · Généré depuis KalanPro AI", small), Spacer(1, 4*mm)]
    for label, value in draft_sections(draft):
        story.append(Paragraph(escape(label), h2))
        if isinstance(value, list):
            items = [ListItem(Paragraph(escape(str(item)), body), leftIndent=8) for item in value]
            story.append(ListFlowable(items, bulletType="bullet", leftIndent=14, bulletColor=BRAND))
        else:
            for paragraph in str(value).split("\n"):
                if paragraph.strip():
                    story.append(Paragraph(escape(paragraph.strip()), body))
    story += [Spacer(1, 5*mm), Paragraph("Relisez toujours un contenu généré par IA avant utilisation.", small)]
    doc.build(story)
    return out.getvalue()
